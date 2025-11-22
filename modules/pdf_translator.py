"""
PDF Translator Module - Full PDF translation with layout preservation
"""
import asyncio
import io
import logging
import os
import uuid
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

logger = logging.getLogger(__name__)

from pdf2zh_next.config.model import SettingsModel
from pdf2zh_next.config.translate_engine_model import GeminiSettings
from pdf2zh_next.high_level import do_translate_async_stream

from config.settings import PDF2ZH_JOBS_ROOT, LANGUAGES, GEMINI_API_KEY
from modules.common import _write_uploaded_file


def _ensure_pdf2zh_job_dir():
    """Create a new job directory for PDF translation."""
    job_dir = PDF2ZH_JOBS_ROOT / str(uuid.uuid4())
    job_dir.mkdir(parents=True, exist_ok=True)
    return job_dir


def _build_pdf2zh_settings(lang_label: str, output_dir: Path):
    """Build settings for pdf2zh translation using Gemini API.
    
    Args:
        lang_label: Target language label
        output_dir: Output directory for translated files
    """
    if not GEMINI_API_KEY:
        raise ValueError("Gemini API key not found. Please set GENAI_API_KEY in your environment.")
    
    # Use gemini-2.0-flash to match other modules (MCQ and Solution generators)
    translate_engine_settings = GeminiSettings(
        gemini_api_key=GEMINI_API_KEY,
        gemini_model="gemini-2.0-flash"
    )
    
    settings = SettingsModel(translate_engine_settings=translate_engine_settings)
    settings.translation.lang_in = "auto"
    settings.translation.lang_out = lang_label
    settings.translation.output = str(output_dir)
    settings.basic.input_files = set()
    # Disable watermark to remove BabelDOC attribution text
    settings.pdf.watermark_output_mode = "no_watermark"
    try:
        settings.validate_settings()
    except Exception as exc:
        logger.warning(f"Settings warning: {exc}")
    return settings


def _run_async(coro):
    """Run async coroutine in a new event loop with proper cleanup."""
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        result = loop.run_until_complete(coro)
        # Give background tasks time to complete
        pending = [t for t in asyncio.all_tasks(loop) if not t.done()]
        if pending:
            # Wait for pending tasks with a timeout
            try:
                loop.run_until_complete(
                    asyncio.wait_for(
                        asyncio.gather(*pending, return_exceptions=True),
                        timeout=2.0
                    )
                )
            except (asyncio.TimeoutError, Exception):
                # Cancel remaining tasks if they don't complete in time
                for task in pending:
                    if not task.done():
                        task.cancel()
                try:
                    loop.run_until_complete(
                        asyncio.gather(*pending, return_exceptions=True)
                    )
                except Exception:
                    pass  # Ignore errors from cancelled tasks
        return result
    finally:
        try:
            # Cancel any remaining tasks
            pending = asyncio.all_tasks(loop)
            for task in pending:
                task.cancel()
            # Give a moment for cleanup
            if pending:
                try:
                    loop.run_until_complete(asyncio.gather(*pending, return_exceptions=True))
                except Exception:
                    pass
        except Exception:
            pass  # Ignore errors during cleanup
        finally:
            try:
                loop.close()
            except Exception:
                pass  # Ignore errors when closing
            asyncio.set_event_loop(None)


async def _stream_pdf2zh(settings: SettingsModel, pdf_path: Path, progress_callback=None):
    """Stream PDF translation events with proper error handling."""
    translate_result = None
    try:
        async for event in do_translate_async_stream(settings, pdf_path):
            event_type = event.get("type")
            if event_type in {"progress_start", "progress_update", "progress_end"}:
                if progress_callback:
                    try:
                        progress_callback(event)
                    except Exception:
                        pass  # Ignore errors in progress callback
            elif event_type == "finish":
                translate_result = event.get("translate_result")
                if translate_result is None:
                    raise RuntimeError("Translation completed but no result was returned.")
                return translate_result
            elif event_type == "error":
                error_msg = event.get("error", "Unknown error")
                error_details = event.get("details", "")
                full_error = f"{error_msg}"
                if error_details:
                    full_error += f" - {error_details}"
                raise RuntimeError(full_error)
    except RuntimeError:
        raise  # Re-raise RuntimeErrors
    except Exception as e:
        # Handle event loop closure gracefully
        if "Event loop is closed" in str(e) or "loop is closed" in str(e).lower():
            # If we have a result, return it despite the error
            if translate_result is not None:
                return translate_result
            # Otherwise, this is a real error
            raise RuntimeError("Translation process was interrupted. Please try again.")
        raise
    
    if translate_result is None:
        raise RuntimeError("Translation stream ended unexpectedly without a result.")
    return translate_result


def translate_pdf_with_pdf2zh(uploaded_file, target_language, status_callback=None):
    """Translate PDF using pdf2zh_next library with Gemini API.
    
    Uses Gemini API as the primary and only translator.
    
    Args:
        uploaded_file: File object or bytes containing PDF data
        target_language: Target language label (e.g., "Hindi", "Telugu")
        status_callback: Optional callback function(status_callback(progress: int, message: str, status: str))
            - progress: 0-100
            - message: Status message
            - status: "processing" | "completed" | "failed"
    """
    lang_code = LANGUAGES.get(target_language, target_language)
    lang_label = target_language
    job_dir = _ensure_pdf2zh_job_dir()
    input_pdf = job_dir / uploaded_file.name
    _write_uploaded_file(uploaded_file, input_pdf)

    def _progress(event, translator_name="Gemini"):
        if status_callback:
            overall_progress = min(max(event.get("overall_progress", 0), 0), 100)
            stage = event.get("stage", "Processing")
            message = f"[{translator_name}] {stage}"
            # Enhanced message with more detail
            if "extract" in stage.lower() or "parse" in stage.lower():
                message = f"📖 PDF Text Extraction: {stage}"
            elif "translat" in stage.lower():
                message = f"🔄 Translation: {stage}"
            elif "render" in stage.lower() or "pdf" in stage.lower() or "save" in stage.lower() or "font" in stage.lower():
                message = f"🎨 PDF Rendering: {stage}"
            else:
                message = f"⚙️ {stage}"
            status_callback(overall_progress, message, "processing")

    # Use Gemini API for translation
    try:
        if status_callback:
            status_callback(0, "🚀 Starting PDF translation process...", "processing")
        
        if not GEMINI_API_KEY:
            raise ValueError("Gemini API key not found. Please set GENAI_API_KEY in your environment.")
        
        settings = _build_pdf2zh_settings(lang_label, job_dir)
        settings.basic.input_files = {str(input_pdf)}
        
        result = _run_async(_stream_pdf2zh(settings, input_pdf, lambda e: _progress(e, "Gemini")))
        
        if result is None:
            raise RuntimeError("Translation completed but returned no result.")
        
        if status_callback:
            status_callback(100, "Translation complete!", "completed")

        return {
            "job_dir": str(job_dir),
            "mono_pdf_path": str(getattr(result, "mono_pdf_path", None)) if result and getattr(result, "mono_pdf_path", None) else None,
            "dual_pdf_path": str(getattr(result, "dual_pdf_path", None)) if result and getattr(result, "dual_pdf_path", None) else None,
            "lang_code": lang_code,
            "lang_label": lang_label,
        }
    
    except Exception as error:
        error_msg = str(error)
        
        # Check for babeldoc-related errors
        if "babeldoc is not available" in error_msg or "babeldoc" in error_msg.lower():
            raise RuntimeError(
                "PDF Translator Feature Unavailable\n\n"
                "The PDF Translator requires the 'babeldoc' library, which cannot be installed due to dependency conflicts.\n\n"
                "Why?\n"
                "babeldoc requires rapidocr-onnxruntime>=1.4.4, but only version 1.2.3 is available on PyPI.\n\n"
                "Current Status:\n"
                "- PDF Translator: NOT WORKING (requires babeldoc)\n"
                "- Solution Generator: WORKING (doesn't need babeldoc)\n"
                "- MCQ Generator: WORKING (doesn't need babeldoc)\n\n"
                "Solutions:\n"
                "1. Wait for rapidocr-onnxruntime>=1.4.4 to be released\n"
                "2. Use Solution Generator or MCQ Generator features which work fine\n"
                "3. Contact babeldoc maintainers about the dependency issue"
            )
        
        # Handle other errors
        if "cannot unpack non-iterable NoneType" in error_msg:
            raise RuntimeError(
                f"Translation failed: The translation service returned an unexpected result. "
                f"This might be due to: 1) Invalid PDF format, 2) Translation service configuration issue, "
                f"3) Missing translation engine settings. Original error: {error_msg}"
            )
        
        raise RuntimeError(
            f"❌ Translation failed:\n\n"
            f"**Error:** {error_msg[:300]}\n\n"
            f"Please check your PDF file and Gemini API key, then try again."
        )


def create_docx_from_pdf(pdf_path: str, title: str):
    """Create DOCX from PDF pages as images."""
    if not pdf_path or not os.path.exists(pdf_path):
        return None
    try:
        doc = Document()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pdf_doc = fitz.open(pdf_path)
        zoom = 200 / 72
        matrix = fitz.Matrix(zoom, zoom)

        for index, page in enumerate(pdf_doc, start=1):
            pix = page.get_pixmap(matrix=matrix)
            image_stream = io.BytesIO(pix.tobytes("png"))
            doc.add_heading(f"Page {index}", level=1)
            doc.add_picture(image_stream, width=doc.sections[0].page_width - Inches(1))
            if index < len(pdf_doc):
                doc.add_page_break()

        pdf_doc.close()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as exc:
        logger.error(f"DOCX creation from PDF failed: {exc}")
        raise RuntimeError(f"Failed to create DOCX from PDF: {exc}") from exc


def create_docx_from_pdf_text(pdf_path: str, title: str):
    """Create DOCX from PDF by extracting text content."""
    if not pdf_path or not os.path.exists(pdf_path):
        return None
    try:
        doc = Document()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        pdf_doc = fitz.open(pdf_path)
        
        for page_num, page in enumerate(pdf_doc, start=1):
            # Extract text from the page
            text = page.get_text()
            
            if text.strip():
                # Add page heading
                doc.add_heading(f"Page {page_num}", level=1)
                
                # Split text into paragraphs and add them
                paragraphs = text.split('\n\n')
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        # Clean up the paragraph
                        para = ' '.join(para.split())  # Normalize whitespace
                        doc.add_paragraph(para)
                
                # Add page break except for last page
                if page_num < len(pdf_doc):
                    doc.add_page_break()
        
        pdf_doc.close()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as exc:
        logger.error(f"DOCX creation from PDF text failed: {exc}")
        raise RuntimeError(f"Failed to create DOCX from PDF text: {exc}") from exc
