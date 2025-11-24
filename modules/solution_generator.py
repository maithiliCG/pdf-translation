"""
solution_generator_final_v3.py

Final stable pipeline (Option A semantics, Simple JSON objects).
- Gemini model: models/gemini-2.5-flash
- No PyMuPDF, no manual extraction, no answer-key scanning
- Solve in English (short explanations up to 3 sentences)
- Translate finalized JSON exactly into target language (no re-solving)
- Build DOCX

Dependencies:
  pip install google-generativeai python-docx
"""

import json
import os
import re
import tempfile
import uuid
from io import BytesIO
from pathlib import Path
from typing import Optional, Union

from config.settings import GEMINI_API_KEY

# ---- external libs ----
try:
    import google.generativeai as genai
except Exception as e:
    raise ImportError("Install google-generativeai: pip install google-generativeai") from e

try:
    from docx import Document
except Exception as e:
    raise ImportError("Install python-docx: pip install python-docx") from e

# ---- config ----
if not GEMINI_API_KEY:
    raise RuntimeError(
        "Gemini API key not configured. Please set GENAI_API_KEY or GEMINI_API_KEY in your environment."
    )

genai.configure(api_key=GEMINI_API_KEY)
MODEL_NAME = os.getenv("GENAI_SOLUTIONS_MODEL") or "models/gemini-2.5-flash"


# ---- helpers ----
def create_docx(text: str, title: str = "Solutions Document") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _upload_file_to_genai(local_path: str):
    """Upload local file to genai and return uploaded file handle/object."""
    return genai.upload_file(local_path)


def _safe_extract_json_array(text: str) -> str:
    """Extract and validate the first JSON array block from model text."""
    text = text.strip()
    # direct parse try
    try:
        parsed = json.loads(text)
        if isinstance(parsed, list):
            return text
    except Exception:
        pass
    # fallback: find first '[' and last ']'
    start = text.find("[")
    end = text.rfind("]")
    if start == -1 or end == -1 or end <= start:
        raise Exception("Model output did not contain a JSON array. Raw (truncated):\n" + text[:1000])
    block = text[start:end+1]
    try:
        parsed = json.loads(block)
    except Exception as exc:
        raise Exception("Extracted JSON array could not be parsed: " + str(exc) + "\nExtracted block (truncated):\n" + block[:1000])
    if not isinstance(parsed, list):
        raise Exception("Extracted JSON is not an array.")
    return block


def _truncate_explanation_to_three_sentences(s: str) -> str:
    """Limit explanation to at most three sentences (naive split)."""
    if not s:
        return s
    sentences = re.split(r'(?<=[\.\?\!])\s+', s.strip())
    if len(sentences) <= 3:
        return s.strip()
    short = " ".join(sentences[:3]).strip()
    if not short.endswith((".", "?", "!")):
        short += "."
    return short


# ---- Gemini calls ----
def _call_gemini_solve_pdf(local_pdf_path: str) -> str:
    """
    Upload PDF and call Gemini to SOLVE in English.
    Return raw model text (expected JSON array).
    """
    uploaded = _upload_file_to_genai(local_pdf_path)
    model = genai.GenerativeModel(MODEL_NAME)

    # NOTE: double braces used where JSON example is inside f-string-like prompt text.
    prompt = (
        "You are an expert MCQ solver. Read the uploaded PDF fully and solve ALL multiple-choice questions (MCQs). "
        "Provide a JSON array of objects. Each object must have keys: number, question, options (array of {\"label\":\"1\",\"text\":\"...\"}), answer (e.g., \"2) 45 litres\"), and explanation.\n\n"
        "IMPORTANT:\n"
        "- RETURN ONLY a valid JSON ARRAY (no explanation outside JSON).\n"
        "- Each explanation MUST be SHORT: 2-3 short sentences (do NOT exceed 3 sentences).\n"
        "- Do NOT add extra commentary, do NOT translate here (English only).\n\n"
        "EXPECTED OBJECT SHAPE (example for one question):\n"
        "[\n"
        "  {{\n"
        "    \"number\": \"1\",\n"
        "    \"question\": \"...\",\n"
        "    \"options\": [{{\"label\":\"1\",\"text\":\"...\"}}, {{\"label\":\"2\",\"text\":\"...\"}}],\n"
        "    \"answer\": \"2) ...\",\n"
        "    \"explanation\": \"Short 2-3 sentence explanation.\"\n"
        "  }}\n"
        "]\n\n"
        "Return only the JSON array."
    )

    response = model.generate_content([uploaded, prompt])
    return getattr(response, "text", str(response))


def _call_gemini_translate_json(json_array_text: str, target_language: str) -> str:
    """
    Translate JSON array text exactly into target_language.
    Preserve structure and numbers. Do NOT re-solve.
    """
    model = genai.GenerativeModel(MODEL_NAME)
    prompt = (
        "Translate the following JSON array into the TARGET_LANGUAGE exactly. ONLY translate natural-language fields "
        "(question, options[*].text, answer text, explanation). Do NOT change numbers, labels, keys, punctuation, or JSON structure. "
        "Do NOT re-solve or modify numeric calculations or answers. RETURN ONLY the translated JSON array.\n\n"
        f"TARGET_LANGUAGE: {target_language}\n\n"
        "JSON INPUT:\n"
        f"{json_array_text}\n\n"
        "Return only the translated JSON array."
    )
    response = model.generate_content([prompt])
    return getattr(response, "text", str(response))


# ---- main pipeline ----
def run_solution_generation_pipeline(
    uploaded_file_or_path: Union[str, object],
    target_language: str = "Telugu",
    progress_callback=None,
    status_callback=None,
    job_id: Optional[str] = None,
):
    """
    uploaded_file_or_path: local path (str) or file-like object with .read() (e.g., Streamlit uploaded file)
    target_language: language name (e.g., 'Telugu','Hindi','Kannada','Tamil','Marathi','Malayalam','Bengali','Odia','Punjabi','Gujarati','English', etc.)
    job_id: optional identifier passed through to status callbacks/logs
    Returns: dict with json_en, json_translated, docx_bytes, language
    """

    # Ensure local file
    if isinstance(uploaded_file_or_path, str):
        local_path = uploaded_file_or_path
    else:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        try:
            tmp.write(uploaded_file_or_path.read())
            tmp.flush()
            local_path = tmp.name
        finally:
            tmp.close()

    def _emit_status(msg: str, frac: float, state: str = "in_progress"):
        frac = 0.0 if frac is None else float(frac)
        frac = max(0.0, min(frac, 1.0))
        pct = int(round(frac * 100))
        formatted_msg = f"[{job_id}] {msg}" if job_id else msg

        if progress_callback:
            try:
                progress_callback(frac)
            except Exception:
                pass

        if status_callback:
            try:
                status_callback(pct, formatted_msg, state)
            except TypeError:
                try:
                    status_callback(formatted_msg)
                except Exception:
                    pass
            except Exception:
                pass

    try:
        _emit_status("1/3: Sending file to Gemini to solve (English)", 0.05)
        raw_solve_text = _call_gemini_solve_pdf(local_path)

        _emit_status("2/3: Parsing solver JSON (English) and enforcing short explanations", 0.4)
        json_block = _safe_extract_json_array(raw_solve_text)
        results_en = json.loads(json_block)

        # Safety: ensure shape and max 3-sentence explanations
        for item in results_en:
            # normalize field names if user model used different keys (try common variants)
            # ensure required keys exist; we'll keep them as-is but attempt to read explanation
            expl = item.get("explanation") or item.get("Explanation") or ""
            item["explanation"] = _truncate_explanation_to_three_sentences(expl)

        _emit_status("3/3: Translating finalized JSON exactly into target language and building DOCX", 0.75)
        json_text_for_translation = json.dumps(results_en, ensure_ascii=False, indent=2)
        translated_raw = _call_gemini_translate_json(json_text_for_translation, target_language)
        translated_block = _safe_extract_json_array(translated_raw)
        results_translated = json.loads(translated_block)

        # Build DOCX
        lines = [f"SOLUTIONS ({target_language})", ""]
        for q in results_translated:
            num = q.get("number", q.get("question_number", ""))
            lines.append(f"Question {num}:")
            qtxt = q.get("question", "")
            if qtxt:
                lines.append(qtxt)
            lines.append("")
            for opt in q.get("options", []):
                lines.append(f"{opt.get('label')}) {opt.get('text')}")
            lines.append("")
            lines.append(f"Answer: {q.get('answer','')}")
            lines.append(f"Explanation: {q.get('explanation','')}")
            lines.append("═══")
            lines.append("")

        docx_bytes = create_docx("\n".join(lines), f"Solutions - {target_language}")

        # Persist DOCX to disk so other modules (Streamlit UI, DB) can read it
        output_dir = Path(tempfile.gettempdir()) / "pdfmath_solutions"
        output_dir.mkdir(parents=True, exist_ok=True)
        docx_filename = f"solutions_{job_id or uuid.uuid4().hex}_{target_language}.docx"
        final_docx_path = output_dir / docx_filename
        with open(final_docx_path, "wb") as docx_file:
            docx_file.write(docx_bytes)

        _emit_status("Done", 1.0, state="completed")

        return {
            "json_en": results_en,
            "json_translated": results_translated,
            "docx_bytes": docx_bytes,
            "language": target_language,
            "job_id": job_id,
            "final_docx": str(final_docx_path),
            "json_data": results_translated,
        }

    except Exception as exc:
        _emit_status(f"Pipeline failed: {exc}", 1.0, state="failed")
        raise

    finally:
        if not isinstance(uploaded_file_or_path, str):
            try:
                os.remove(local_path)
            except Exception:
                pass


# ---- __main__ example (developer-provided file path from conversation history) ----
if __name__ == "__main__":
    # Developer note: using the uploaded file path from conversation history as example.
    sample_path = "/mnt/data/solutions_Telugu (9).docx"  # path from conversation history; change to a PDF if needed
    if not os.path.exists(sample_path):
        print("Example file not found:", sample_path)
    else:
        print("Running pipeline on:", sample_path)
        out = run_solution_generation_pipeline(sample_path, target_language="Telugu", progress_callback=print, status_callback=print)
        print("Sample English item:", out["json_en"][0] if out["json_en"] else None)
        print("Sample Translated item:", out["json_translated"][0] if out["json_translated"] else None)
        with open(f"solutions_{out['language']}.docx", "wb") as f:
            f.write(out["docx_bytes"])
        print("Wrote:", f"solutions_{out['language']}.docx")
