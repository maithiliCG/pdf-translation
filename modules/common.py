"""
Common utilities and shared functions for PDFMathTranslate
"""
import html
import io
import json
import logging
import re
import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from config.settings import model

logger = logging.getLogger(__name__)


def _call_generative_model(prompt: str, max_attempts: int = 3, cooldown_seconds: float = 45.0):
    """
    Wrapper around Gemini calls with basic retry/backoff on quota errors.
    """
    if model is None:
        raise RuntimeError("Gemini API key not configured. Please set GENAI_API_KEY in your environment.")
    
    last_error: Exception | None = None
    for attempt in range(1, max_attempts + 1):
        try:
            response = model.generate_content(prompt)
            time.sleep(1.0)  # keep requests under per-minute free-tier limits
            return response
        except Exception as exc:
            last_error = exc
            message = str(exc).lower()
            
            # Check for quota/rate limit errors
            if "429" in message or "quota" in message or "rate limit" in message or "resourceexhausted" in message:
                error_msg = str(exc)
                # Extract retry delay if available
                retry_delay = 60  # default
                if "retry in" in error_msg.lower():
                    delay_match = re.search(r"retry in ([\d.]+)s", error_msg.lower())
                    if delay_match:
                        retry_delay = int(float(delay_match.group(1))) + 5
                
                if attempt < max_attempts:
                    # Silently retry without showing warnings to user
                    time.sleep(cooldown_seconds)
                    continue
                else:
                    # Final attempt failed - raise error but don't show quota message
                    raise RuntimeError(
                        f"API request failed after {max_attempts} attempts. Please try again later."
                    )
            break
    
    # If we get here, it's not a quota error
    if last_error:
        raise last_error
    raise RuntimeError("Gemini call failed unexpectedly.")


def extract_json_block(text: str) -> str:
    """Extract JSON block from text."""
    if not text:
        return ""
    match = re.search(r"```json\s*(.*?)```", text, re.DOTALL | re.IGNORECASE)
    if match:
        return match.group(1).strip()
    match = re.search(r"```\s*(.*?)```", text, re.DOTALL)
    if match:
        return match.group(1).strip()
    match = re.search(r"\[\s*\{.*?\}\s*\]", text, re.DOTALL)
    if match:
        return match.group(0).strip()
    return text.strip()


def extract_inner_json(text: str):
    """Extract and parse JSON from text."""
    if not text:
        return None
    match = re.search(r"```json\s*(.*?)```", text, re.DOTALL | re.IGNORECASE)
    if not match:
        return None
    try:
        return json.loads(match.group(1))
    except json.JSONDecodeError:
        return None


def _clean_text(value: str) -> str:
    """Clean HTML entities and whitespace from text."""
    if not value:
        return ""
    value = html.unescape(str(value))
    value = (
        value.replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&amp;", "&")
        .strip()
    )
    return value


def _write_uploaded_file(uploaded_file, destination: Path):
    """Write uploaded file to destination path."""
    with open(destination, "wb") as f:
        if hasattr(uploaded_file, "getvalue"):
            f.write(uploaded_file.getvalue())
        elif isinstance(uploaded_file, bytes):
            f.write(uploaded_file)
        else:
            f.write(uploaded_file.read())


def create_docx(content, title="Document"):
    """Create a DOCX document from text content."""
    try:
        doc = Document()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in content.split("\n"):
            text = line.strip()
            if not text:
                doc.add_paragraph("")
                continue
            if text.startswith("**Question"):
                p = doc.add_heading(text.replace("**", "").replace(":", ""), level=2)
                p.runs[0].font.color.rgb = RGBColor(0, 0, 255)
            elif text.startswith("**Correct Answer"):
                p = doc.add_heading("Correct Answer", level=3)
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            elif text.startswith("**Explanation"):
                p = doc.add_heading("Explanation", level=3)
                p.runs[0].font.color.rgb = RGBColor(128, 0, 128)
            elif text.startswith("═══"):
                doc.add_paragraph("_" * 60)
            elif text.strip().startswith("✓"):
                # Correct option - make it bold and green
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif text.strip().startswith("  ") and (")" in text or ")" in text):
                # Regular option - indent it
                p = doc.add_paragraph(text.strip(), style="List Bullet")
            else:
                doc.add_paragraph(text)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as exc:
        logger.error(f"DOCX creation failed: {exc}")
        raise RuntimeError(f"Failed to create DOCX document: {exc}")

